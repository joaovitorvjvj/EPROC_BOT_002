import os
from docx import Document
from app.core.config import settings
from app.core.logger import logger

class DocxService:
    def __init__(self):
        self.output_dir = os.path.join(settings.STORAGE_PATH, "docx")

    async def generate_document(self, process_data: dict) -> str:
        try:
            doc = Document()
            doc.add_heading(f'Relatório: {process_data.get("name")}', 0)
            canvas = process_data.get("canvas_data", {})
            
            for key, val in canvas.items():
                doc.add_heading(key.replace("_", " ").title(), level=1)
                doc.add_paragraph(str(val))

            file_path = os.path.join(self.output_dir, f"{process_data['id']}.docx")
            doc.save(file_path)
            return file_path
        except Exception as e:
            logger.error(f"Erro DOCX: {str(e)}")
            raise e